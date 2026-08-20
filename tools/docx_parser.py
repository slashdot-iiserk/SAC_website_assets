#!/usr/bin/env python3
"""Parse DOCX files to Markdown, extracting text and embedded images with context."""

import os
import re
import sys
from pathlib import Path
from PIL import Image
from docx import Document
from docx.oxml.ns import qn


WEBP_QUALITY = 85


def extract_images_with_context(
    doc: Document, docx_path: Path, output_dir: Path
) -> list[dict]:
    """Extract images from DOCX with surrounding text context for captions."""
    images = []
    output_dir.mkdir(parents=True, exist_ok=True)

    img_counter = 0

    for para in doc.paragraphs:
        para_text = para.text.strip()

        for run in para.runs:
            drawing_elements = run._element.findall(f".//{qn('w:drawing')}")

            for drawing in drawing_elements:
                blip_elements = drawing.findall(f".//{qn('a:blip')}")

                for blip in blip_elements:
                    embed_id = blip.get(qn("r:embed"))
                    if not embed_id:
                        continue

                    try:
                        rel = doc.part.rels.get(embed_id)
                        if not rel or "image" not in rel.reltype:
                            continue

                        img_data = rel.target_part.blob
                        ext = Path(rel.target_part.partname).suffix or ".png"

                        img_counter += 1
                        img_name = f"img_{img_counter:03d}{ext}"
                        img_path = output_dir / img_name
                        img_path.write_bytes(img_data)

                        caption = para_text if para_text else f"Image {img_counter}"

                        images.append(
                            {
                                "rId": embed_id,
                                "path": img_path,
                                "name": img_name,
                                "caption": caption,
                                "context": para_text,
                            }
                        )
                    except Exception:
                        continue

    for rel_id, rel in doc.part.rels.items():
        if "image" not in rel.reltype:
            continue

        already_extracted = any(img["rId"] == rel_id for img in images)
        if already_extracted:
            continue

        try:
            img_data = rel.target_part.blob
            ext = Path(rel.target_part.partname).suffix or ".png"

            img_counter += 1
            img_name = f"img_{img_counter:03d}{ext}"
            img_path = output_dir / img_name
            img_path.write_bytes(img_data)

            images.append(
                {
                    "rId": rel_id,
                    "path": img_path,
                    "name": img_name,
                    "caption": f"Image {img_counter}",
                    "context": "",
                }
            )
        except Exception:
            continue

    return images


def convert_image_to_webp(src: Path) -> Path:
    """Convert image to WebP format, return new path."""
    if src.suffix.lower() == ".webp":
        return src

    dst = src.with_suffix(".webp")
    if dst.exists():
        return dst

    try:
        with Image.open(src) as img:
            if img.mode == "RGBA":
                img = img.convert("RGB")
            elif img.mode not in ("RGB", "L"):
                img = img.convert("RGB")

            w, h = img.size
            if max(w, h) > 2000:
                ratio = 2000 / max(w, h)
                img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)

            img.save(dst, "WEBP", quality=WEBP_QUALITY, method=6)
            src.unlink()
            return dst
    except Exception as e:
        print(f"    WARN: Could not convert {src.name} to WebP: {e}", file=sys.stderr)
        return src


def docx_to_markdown(
    docx_path: Path,
    output_dir: Path,
    output_stem: str | None = None,
    image_dir_name: str | None = None,
) -> Path:
    """Convert DOCX to Markdown with embedded images and captions."""
    doc = Document(str(docx_path))
    output_stem = output_stem or docx_path.stem
    img_dir = output_dir / (image_dir_name or f"{output_stem}_images")
    images = extract_images_with_context(doc, docx_path, img_dir)

    md_lines = []
    md_lines.append(f"# {docx_path.stem.replace('_', ' ').replace('-', ' ').title()}\n")

    img_lookup = {img["rId"]: img for img in images}
    img_counter = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        para_has_image = False
        for run in para.runs:
            drawings = run._element.findall(f".//{qn('w:drawing')}")
            for drawing in drawings:
                blips = drawing.findall(f".//{qn('a:blip')}")
                for blip in blips:
                    embed_id = blip.get(qn("r:embed"))
                    if embed_id and embed_id in img_lookup:
                        para_has_image = True
                        img_info = img_lookup[embed_id]

                        webp_path = convert_image_to_webp(img_info["path"])
                        img_counter += 1
                        image_ref = webp_path.relative_to(output_dir).as_posix()

                        caption = img_info["caption"]
                        if caption and caption != f"Image {img_counter}":
                            caption = caption[:100]
                            md_lines.append(f"\n![{caption}]({image_ref})\n")
                        else:
                            md_lines.append(f"\n![Image {img_counter}]({image_ref})\n")

        if para_has_image:
            continue

        if not text:
            md_lines.append("")
            continue

        style = para.style.name if para.style else ""

        if "Heading 1" in style:
            md_lines.append(f"# {text}\n")
        elif "Heading 2" in style:
            md_lines.append(f"## {text}\n")
        elif "Heading 3" in style:
            md_lines.append(f"### {text}\n")
        elif "List" in style:
            md_lines.append(f"- {text}")
        else:
            md_lines.append(f"{text}\n")

    for table in doc.tables:
        if not table.rows:
            continue
        md_lines.append(
            "\n| "
            + " | ".join(cell.text.strip() for cell in table.rows[0].cells)
            + " |"
        )
        md_lines.append("| " + " | ".join("---" for _ in table.rows[0].cells) + " |")
        for row in table.rows[1:]:
            md_lines.append(
                "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
            )
        md_lines.append("")

    remaining_imgs = [
        img
        for img in images
        if not any(
            blip.get(qn("r:embed")) == img["rId"]
            for para in doc.paragraphs
            for run in para.runs
            for drawing in run._element.findall(f".//{qn('w:drawing')}")
            for blip in drawing.findall(f".//{qn('a:blip')}")
        )
    ]

    for img in remaining_imgs:
        webp_path = convert_image_to_webp(img["path"])
        image_ref = webp_path.relative_to(output_dir).as_posix()
        md_lines.append(f"\n![{img['caption']}]({image_ref})\n")

    md_path = output_dir / f"{output_stem}.md"
    md_path.write_text("\n".join(md_lines), encoding="utf-8")
    return md_path


def process_directory(src_dir: Path, dst_dir: Path) -> dict:
    """Process all DOCX files in directory."""
    stats = {"processed": 0, "skipped": 0, "errors": 0}

    for root, _, files in os.walk(src_dir):
        for fname in files:
            src = Path(root) / fname
            if src.suffix.lower() not in (".docx",):
                continue

            rel = src.relative_to(src_dir)
            out = dst_dir / rel.parent
            out.mkdir(parents=True, exist_ok=True)

            try:
                docx_to_markdown(src, out)
                stats["processed"] += 1
                print(f"  OK: {rel}")
            except Exception as e:
                stats["errors"] += 1
                print(f"  ERROR: {rel}: {e}", file=sys.stderr)

    return stats


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Convert DOCX to Markdown")
    parser.add_argument("source", help="Source directory")
    parser.add_argument("-o", "--output", help="Output directory (default: source)")
    args = parser.parse_args()

    src = Path(args.source).resolve()
    dst = Path(args.output).resolve() if args.output else src

    if not src.exists():
        print(f"Source not found: {src}", file=sys.stderr)
        sys.exit(1)

    print(f"Processing DOCX: {src} -> {dst}")
    stats = process_directory(src, dst)
    print(
        f"\nDone: {stats['processed']} processed, {stats['skipped']} skipped, {stats['errors']} errors"
    )


if __name__ == "__main__":
    main()
